import html
import io
import json
import os
import re
from pathlib import Path
from dataclasses import dataclass

import streamlit as st
import streamlit.components.v1 as components
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_COLOR_INDEX
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

try:
    from openai import OpenAI
except Exception:  # pragma: no cover - optional dependency at runtime
    OpenAI = None


st.set_page_config(page_title="单词造句卡片", page_icon="A", layout="wide")

LOGO_PATH = Path(__file__).parent / "assets" / "zhejiang_university_logo.png"
INSTITUTE_LOGO_PATH = Path(__file__).parent / "assets" / "shanghai_veterinary_research_institute.png"


@dataclass
class Card:
    word: str
    level: str
    sentence: str
    meaning_zh: str
    sentence_zh: str
    extra_sentence: str = ""
    extra_sentence_zh: str = ""


@dataclass
class AIConfig:
    provider: str
    api_key: str
    model: str
    base_url: str


PROVIDER_PRESETS = {
    "OpenAI": {
        "secret": "OPENAI_API_KEY",
        "model": "gpt-4.1-mini",
        "base_url": "",
    },
    "DeepSeek": {
        "secret": "DEEPSEEK_API_KEY",
        "model": "deepseek-chat",
        "base_url": "https://api.deepseek.com",
    },
    "通义千问": {
        "secret": "DASHSCOPE_API_KEY",
        "model": "qwen-plus",
        "base_url": "https://dashscope.aliyuncs.com/compatible-mode/v1",
    },
    "Kimi": {
        "secret": "MOONSHOT_API_KEY",
        "model": "moonshot-v1-8k",
        "base_url": "https://api.moonshot.cn/v1",
    },
    "自定义": {
        "secret": "AI_API_KEY",
        "model": "gpt-4.1-mini",
        "base_url": "",
    },
}


FALLBACK_MEANINGS = {
    "watch tv": "看电视",
    "read books": "看书",
    "play chess": "下棋",
    "keep pets": "养宠物",
    "grow flowers": "种花",
    "walk": "走路/散步",
    "walking": "走路/散步",
    "scroll through mobile phones": "刷手机/浏览手机",
    "important": "重要的",
    "beautiful": "美丽的",
    "practice": "练习",
    "remember": "记住",
    "improve": "提高/改善",
    "friendly": "友好的",
    "careful": "小心的/仔细的",
    "quickly": "快速地",
}


def get_secret(name: str) -> str:
    try:
        return st.secrets.get(name, "")
    except Exception:
        return os.getenv(name, "")


def parse_words(raw: str) -> list[tuple[str, str]]:
    items: list[tuple[str, str]] = []
    for line in raw.replace("，", ",").splitlines():
        parts = [p.strip() for p in line.split(",") if p.strip()]
        for part in parts:
            if "|" in part:
                word, meaning = [x.strip() for x in part.split("|", 1)]
            elif "=" in part:
                word, meaning = [x.strip() for x in part.split("=", 1)]
            else:
                word, meaning = part, ""
            if word:
                items.append((word, meaning))
    return items


def ai_cards(words: list[tuple[str, str]], level: str, config: AIConfig) -> list[Card] | None:
    if not config.api_key or not config.model or OpenAI is None:
        return None

    client_kwargs = {"api_key": config.api_key}
    if config.base_url:
        client_kwargs["base_url"] = config.base_url.rstrip("/")
    client = OpenAI(**client_kwargs)

    word_payload = [{"word": word, "meaning_hint": meaning} for word, meaning in words]
    prompt = f"""
请为英语学习者生成单词卡片。难度约为 {level}，但不要在结果里显示等级。
输入是 JSON 数组，每项包含 word 和可选 meaning_hint。
请返回严格 JSON 数组，不要 Markdown。每项字段：
word: 原词或短语
level: 固定返回空字符串
sentence: 一个自然、简短、适合朗读的英文例句，必须包含原词或短语
meaning_zh: 中文意思
sentence_zh: 英文例句的中文翻译
extra_sentence: 另一个自然英文例句，也必须包含原词或短语，不要和 sentence 重复
extra_sentence_zh: extra_sentence 的中文翻译

要求：
- 不要使用固定模板句。
- 每个例句要像真实生活中会说的话。
- 如果 meaning_hint 不为空，中文意思必须优先使用或贴近 meaning_hint。
- sentence 和 extra_sentence 必须直接包含输入的 word 或短语，方便之后高亮。

输入：
{json.dumps(word_payload, ensure_ascii=False)}
"""
    try:
        response = client.chat.completions.create(
            model=config.model,
            messages=[
                {"role": "system", "content": "You write accurate bilingual English vocabulary flashcards."},
                {"role": "user", "content": prompt},
            ],
            temperature=0.4,
        )
    except Exception as exc:
        st.warning(f"{config.provider} 调用失败，已改用离线模板。错误：{exc}")
        return None
    text = response.choices[0].message.content or "[]"
    try:
        data = json.loads(text)
    except json.JSONDecodeError:
        match = re.search(r"\[.*\]", text, re.S)
        data = json.loads(match.group(0)) if match else []

    cards: list[Card] = []
    for item in data:
        cards.append(
            Card(
                word=str(item.get("word", "")).strip(),
                level=str(item.get("level", level)).strip() or level,
                sentence=str(item.get("sentence", "")).strip(),
                meaning_zh=str(item.get("meaning_zh", "")).strip(),
                sentence_zh=str(item.get("sentence_zh", "")).strip(),
                extra_sentence=str(item.get("extra_sentence", "")).strip(),
                extra_sentence_zh=str(item.get("extra_sentence_zh", "")).strip(),
            )
        )
    return [card for card in cards if card.word and card.sentence]


def fallback_cards(words: list[tuple[str, str]], level: str) -> list[Card]:
    cards: list[Card] = []
    for word, meaning_hint in words:
        key = word.lower()
        meaning = meaning_hint or FALLBACK_MEANINGS.get(key, "请填写中文意思")
        sentence = make_fallback_sentence(word)
        cards.append(
            Card(
                word=word,
                level=level,
                sentence=sentence,
                meaning_zh=meaning,
                sentence_zh="请根据例句补充中文翻译",
                extra_sentence=make_extra_fallback_sentence(word),
                extra_sentence_zh="请根据额外例句补充中文翻译",
            )
        )
    return cards


def make_fallback_sentence(word: str) -> str:
    clean = word.strip()
    lower = clean.lower()
    if " " in clean:
        return f"I try to {clean} every day after school."
    if lower.endswith("ly"):
        return f"She answered the question {clean} in class."
    if lower.endswith("ing"):
        return f"{clean.capitalize()} helps me relax after a busy day."
    return f"I want to use the word {clean} in a clear English sentence."


def make_extra_fallback_sentence(word: str) -> str:
    clean = word.strip()
    if " " in clean:
        return f"My teacher asked me to {clean} in a short story."
    return f"Please remember how to use {clean} when you speak English."


def make_cards(words: list[tuple[str, str]], level: str, config: AIConfig) -> list[Card]:
    generated = ai_cards(words, level, config)
    return generated if generated else fallback_cards(words, level)


def split_highlight(text: str, target: str) -> list[tuple[str, bool]]:
    tokens = [target.strip()]
    if " / " in target:
        tokens.extend(part.strip() for part in target.split("/") if part.strip())
    tokens = sorted(set(filter(None, tokens)), key=len, reverse=True)
    if not tokens:
        return [(text, False)]
    pattern = re.compile("(" + "|".join(re.escape(t) for t in tokens) + ")", re.I)
    parts: list[tuple[str, bool]] = []
    last = 0
    for match in pattern.finditer(text):
        if match.start() > last:
            parts.append((text[last : match.start()], False))
        parts.append((match.group(0), True))
        last = match.end()
    if last < len(text):
        parts.append((text[last:], False))
    return parts


def highlighted_html(text: str, target: str) -> str:
    result = []
    for part, is_hit in split_highlight(text, target):
        safe = html.escape(part)
        if is_hit:
            result.append(f'<span style="color:#c00000;font-weight:700;">{safe}</span>')
        else:
            result.append(safe)
    return "".join(result)


def highlighted_extra_html(text: str, target: str, meaning_zh: str) -> str:
    result = []
    annotated = False
    meaning = html.escape(meaning_zh.strip())
    for part, is_hit in split_highlight(text, target):
        safe = html.escape(part)
        if is_hit:
            result.append(f'<span style="color:#c00000;font-weight:700;">{safe}</span>')
            if meaning and not annotated:
                result.append(f" ({meaning})")
                annotated = True
        else:
            result.append(safe)
    return "".join(result)


def set_run_font(run, ascii_font: str, east_asia_font: str | None = None) -> None:
    east_asia_font = east_asia_font or ascii_font
    run.font.name = ascii_font
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), ascii_font)
    r_fonts.set(qn("w:hAnsi"), ascii_font)
    r_fonts.set(qn("w:eastAsia"), east_asia_font)


def add_highlighted_run(paragraph, text: str, target: str) -> None:
    for part, is_hit in split_highlight(text, target):
        run = paragraph.add_run(part)
        set_run_font(run, "Cambria")
        if is_hit:
            run.font.color.rgb = RGBColor(192, 0, 0)
            run.font.bold = True
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def add_red_run(paragraph, text: str) -> None:
    run = paragraph.add_run(text)
    set_run_font(run, "Microsoft YaHei", "微软雅黑")
    run.font.color.rgb = RGBColor(192, 0, 0)
    run.font.bold = True


def set_cell_margins(cell, top: int = 160, start: int = 220, bottom: int = 100, end: int = 220) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        margin = tc_mar.find(qn(f"w:{margin_name}"))
        if margin is None:
            margin = OxmlElement(f"w:{margin_name}")
            tc_mar.append(margin)
        margin.set(qn("w:w"), str(value))
        margin.set(qn("w:type"), "dxa")


def build_docx(cards: list[Card]) -> bytes:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    table.columns[0].width = Inches(2.45)
    table.columns[1].width = Inches(5.4)

    headers = ["Front (Phrase)", "Back (Sentence + Meaning)"]
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_margins(cell, top=160, start=220, bottom=160, end=220)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(text)
        set_run_font(run, "Microsoft YaHei", "微软雅黑")
        run.bold = True
        run.font.size = Pt(10.5)

    for card in cards:
        cells = table.add_row().cells
        front = cells[0].paragraphs[0]
        front_run = front.add_run(card.word)
        set_run_font(front_run, "Cambria")

        back = cells[1].paragraphs[0]

        add_highlighted_run(back, card.sentence, card.word)
        back.add_run("\n")
        add_red_run(back, card.meaning_zh)
        back.add_run("\n")
        translation_run = back.add_run(card.sentence_zh)
        set_run_font(translation_run, "Microsoft YaHei", "微软雅黑")

        for cell in cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.line_spacing = 1.35
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.size = Pt(10.5)

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


def speech_component(card: Card, index: int) -> None:
    sentence = card.sentence
    word_to_speak = card.word.split("/")[0].strip()
    markup = highlighted_html(sentence, card.word)
    js_sentence = json.dumps(sentence)
    js_word = json.dumps(word_to_speak)
    components.html(
        f"""
        <style>
          .sentence-box {{
            font-family: Arial, sans-serif;
            font-size: 18px;
            line-height: 1.55;
            color: #1f2937;
            padding: 12px 0;
          }}
          .target-word {{
            color: #d11f1f;
            font-weight: 700;
          }}
          .speech-controls {{
            align-items: center;
            display: flex;
            gap: 8px;
          }}
          .speak-btn {{
            border: 1px solid #d0d7de;
            background: white;
            width: 42px;
            height: 42px;
            cursor: pointer;
            font-size: 20px;
            line-height: 1;
          }}
          .speak-sentence {{
            border-radius: 999px;
          }}
          .speak-word {{
            border-radius: 8px;
          }}
          .speak-btn:hover {{ background: #f6f8fa; }}
          .speech-error {{
            color: #8a6d00;
            display: inline-block;
            font-family: Arial, sans-serif;
            font-size: 12px;
            margin-left: 8px;
          }}
        </style>
        <div class="sentence-box">{markup}</div>
        <div class="speech-controls">
          <button class="speak-btn speak-sentence" title="播放句子" aria-label="播放句子" onclick='speak_{index}({js_sentence})'>🔊</button>
          <button class="speak-btn speak-word" title="播放单词" aria-label="播放单词" onclick='speak_{index}({js_word})'>♪</button>
          <span id="speech-error-{index}" class="speech-error"></span>
        </div>
        <script>
          function pickEnglishVoice_{index}() {{
            const voices = window.speechSynthesis.getVoices();
            return voices.find(v => v.lang === "en-US")
              || voices.find(v => v.lang && v.lang.startsWith("en"))
              || null;
          }}

          function speak_{index}(text) {{
            const errorBox = document.getElementById("speech-error-{index}");
            errorBox.textContent = "";
            if (!("speechSynthesis" in window)) {{
              errorBox.textContent = "当前浏览器不支持朗读";
              return;
            }}

            const utterance = new SpeechSynthesisUtterance(text);
            utterance.lang = "en-US";
            utterance.rate = 0.88;
            utterance.pitch = 1;
            utterance.volume = 1;
            const voice = pickEnglishVoice_{index}();
            if (voice) {{
              utterance.voice = voice;
            }}
            utterance.onerror = () => {{
              errorBox.textContent = "朗读失败，请检查浏览器/系统声音";
            }};

            window.speechSynthesis.cancel();
            window.speechSynthesis.resume();
            window.speechSynthesis.speak(utterance);
          }}

          window.speechSynthesis.onvoiceschanged = pickEnglishVoice_{index};
        </script>
        """,
        height=115,
    )


title_col, zju_col, institute_col = st.columns([3.2, 1.1, 3.4], vertical_alignment="center")
with title_col:
    st.title("单词造句卡片")
    st.caption("输入单词或短语，生成可下载的 Word 卡片；也可以直接播放英文例句并高亮目标词。")
with zju_col:
    if LOGO_PATH.exists():
        st.image(str(LOGO_PATH), width=165)
with institute_col:
    if INSTITUTE_LOGO_PATH.exists():
        st.image(str(INSTITUTE_LOGO_PATH), use_container_width=True)

with st.sidebar:
    st.header("设置")
    provider_options = list(PROVIDER_PRESETS.keys())
    provider = st.selectbox("AI 引擎", provider_options, index=provider_options.index("DeepSeek"))
    preset = PROVIDER_PRESETS[provider]
    default_api_key = get_secret(preset["secret"]) or get_secret("AI_API_KEY")
    api_key_input = st.text_input(
        "API key",
        value="",
        type="password",
        placeholder=f"可留空，优先读取 Secrets: {preset['secret']}",
    )
    model = st.text_input("模型名", value=get_secret("AI_MODEL") or preset["model"])
    base_url = st.text_input("Base URL", value=get_secret("AI_BASE_URL") or preset["base_url"])
    st.write("可选格式：`word | 中文意思`，没有 API key 时会用中文提示或内置小词库。")
    st.write("部署到 Streamlit 后，在 Secrets 里添加 `DEEPSEEK_API_KEY` 即可默认使用 DeepSeek。")

level = "A2"
ai_config = AIConfig(
    provider=provider,
    api_key=api_key_input.strip() or default_api_key,
    model=model.strip(),
    base_url=base_url.strip(),
)

default_words = "scroll through mobile phones | 刷手机/浏览手机\nwatch TV | 看电视\nwalk / walking | 散步/走路"
raw_words = st.text_area("输入单词或短语", value=default_words, height=160)

col_a, col_b = st.columns([1, 1])
with col_a:
    generate = st.button("生成卡片", type="primary", use_container_width=True)
with col_b:
    clear = st.button("清空", use_container_width=True)

if clear:
    st.session_state.pop("cards", None)

if generate:
    parsed = parse_words(raw_words)
    if not parsed:
        st.warning("请先输入至少一个单词或短语。")
    else:
        if not ai_config.api_key:
            st.info("没有检测到 API key，先使用离线模板生成。填入 API key 后会联网调用 AI 生成更自然的句子。")
        with st.spinner("正在联网生成例句和 Word..."):
            st.session_state.cards = make_cards(parsed, level, ai_config)

cards: list[Card] = st.session_state.get("cards", [])
if cards:
    docx_bytes = build_docx(cards)
    st.download_button(
        "下载 Word 文件",
        data=docx_bytes,
        file_name="flashcards_generated.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True,
    )

    for idx, card in enumerate(cards):
        with st.container(border=True):
            st.subheader(card.word)
            speech_component(card, idx)
            safe_meaning = html.escape(card.meaning_zh)
            safe_translation = html.escape(card.sentence_zh)
            st.markdown(
                f'<div style="color:#c00000;font-weight:700;font-size:18px;margin-top:8px;">{safe_meaning}</div>',
                unsafe_allow_html=True,
            )
            st.markdown(
                f'<div style="font-size:18px;margin-top:6px;">{safe_translation}</div>',
                unsafe_allow_html=True,
            )
            extra_sentence = getattr(card, "extra_sentence", "")
            if extra_sentence:
                with st.expander("＋", expanded=False):
                    extra_markup = highlighted_extra_html(extra_sentence, card.word, card.meaning_zh)
                    st.markdown(
                        f'<div style="font-size:18px;line-height:1.55;">{extra_markup}</div>',
                        unsafe_allow_html=True,
                    )
                    extra_translation = getattr(card, "extra_sentence_zh", "")
                    if extra_translation:
                        st.markdown(
                            f'<div style="font-size:17px;margin-top:6px;">{html.escape(extra_translation)}</div>',
                            unsafe_allow_html=True,
                        )
